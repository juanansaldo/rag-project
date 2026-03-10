from pathlib import Path

from app.loaders import (
    load_text_file, 
    load_document, 
    load_directory,
    load_html,
    load_csv,
    load_docx,
)


def test_load_text_file(tmp_path):
    f = tmp_path / "sample.txt"
    f.write_text("Hello from test file.")
    result = load_text_file(f)
    assert len(result) == 1
    assert result[0]["text"] == "Hello from test file."
    assert result[0]["source"] == "sample.txt"
    assert result[0]["page"] == 0


def test_load_text_file_missing_returns_empty():
    assert load_text_file(Path("/nonexistent/path.txt")) == []


def test_load_document_txt(tmp_path):
    f = tmp_path / "doc.txt"
    f.write_text("Content here.")
    result = load_document(f)
    assert len(result) == 1
    assert "Content here." in result[0]["text"]
    assert result[0]["source"] == "doc.txt"


def test_load_document_md(tmp_path):
    f = tmp_path / "readme.md"
    f.write_text("# Title\n\nSome markdown.")
    result = load_document(f)
    assert len(result) == 1
    assert "Title" in result[0]["text"]


def test_load_directory(tmp_path):
    (tmp_path / "a.txt").write_text("File A")
    (tmp_path / "b.txt").write_text("File B")
    result = load_directory(tmp_path)
    assert len(result) == 2
    texts = [r["text"] for r in result]
    assert "File A" in texts
    assert "File B" in texts


def test_load_directory_nonexistent_returns():
    assert load_directory(Path("/nonexistent/dir")) == []


def test_load_document_html(tmp_path):
    f = tmp_path / "page.html"
    f.write_text("<html><body><p>Hello</p> <p>World</p></body></html>", encoding="utf-8")
    result = load_document(f)
    assert len(result) == 1
    assert "Hello" in result[0]["text"]
    assert "World" in result[0]["text"]
    assert result[0]["source"] == "page.html"
    assert result[0]["page"] == 0


def test_load_html_strips_tags(tmp_path):
    f = tmp_path / "a.html"
    f.write_text("<div>Only this counts</div>", encoding="utf-8")
    result = load_html(f)
    assert len(result) == 1
    assert "Only this counts" in result[0]["text"]


def test_load_document_csv(tmp_path):
    f = tmp_path / "data.csv"
    f.write_text("a,b,c\n1,2,3\n", encoding="utf-8")
    result = load_document(f)
    assert len(result) == 1
    assert "a" in result[0]["text"] and "1" in result[0]["text"]
    assert result[0]["source"] == "data.csv"


def test_load_csv_empty_returns_empty(tmp_path):
    f = tmp_path / "empty.csv"
    f.write_text("", encoding="utf-8")
    result = load_csv(f)
    assert result == []


def test_load_directory_includes_html_and_csv(tmp_path):
    (tmp_path / "a.txt").write_text("A")
    (tmp_path / "b.html").write_text("<p>B</p>")
    (tmp_path / "c.csv").write_text("x,y\n1,2")
    result = load_directory(tmp_path)
    assert len(result) == 3
    texts = [r["text"] for r in result]
    assert any("A" in t for t in texts)
    assert any("B" in t for t in texts)
    assert any("x" in t and "1" in t for t in texts)


def test_load_document_docx(tmp_path):
    """Load DOCX via load_document returns paragraph text"""
    try:
        from docx import Document
    except ImportError:
        import pytest
        pytest.skip("python-docx not installed")
    doc_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.save(str(doc_path))
    result = load_document(doc_path)
    assert len(result) == 1
    assert "First paragraph." in result[0]["text"]
    assert "Second paragraph." in result[0]["text"]
    assert result[0]["source"] == "sample.docx"
    assert result[0]["page"] == 0